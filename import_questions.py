"""
题库自动导入管线
================
将 PDF / 文本文件自动解析为结构化题目，分类到对应知识点，评估难度，
写入 data/imported_exercises.json。

knowledge_base.py 在加载时会自动合并该文件，无需手动修改代码。

用法:
    python import_questions.py                          # 扫描 data/ 下所有 OCR/提取文本
    python import_questions.py data/GZMK数学试卷2_ocr.txt  # 导入指定文件
    python import_questions.py --from-pdf "试卷.pdf"    # 先提取再导入
    python import_questions.py --from-docx "题库.docx"  # 从 Word 文档导入
"""

import json
import os
import re
import sys
import hashlib
from typing import Optional

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
IMPORT_FILE = os.path.join(DATA_DIR, "imported_exercises.json")

# ---------------------------------------------------------------------------
# 知识点关键词映射表 —— 用于自动分类
# ---------------------------------------------------------------------------
# 格式: topic_name -> [关键词列表]  (命中越多越优先)
TOPIC_KEYWORDS = {
    # ===== 1.1 数的认识 =====
    "数的分类": ["自然数", "整数", "分数", "小数", "百分数", "互化", "正数", "负数有"],
    "数的组成（大数读写改写）": ["读作", "写作", "改写", "省略", "四舍五入", "亿", "万"],
    "分数的意义、分类与性质": ["真分数", "假分数", "带分数", "分数的基本性质", "通分", "约分"],
    "用分数表示数量和占比": ["占", "几分之几", "分率", "单位1"],
    "负数": ["负数", "正数", "零下", "海拔", "相反数"],

    # ===== 1.2 因数与倍数 =====
    "因数与倍数": ["因数", "倍数", "公因数", "公倍数", "最大公因数", "最小公倍数"],
    "2、3、5倍数的特征": ["偶数", "奇数", "能被2整除", "能被3整除", "能被5整除"],
    "质数与合数": ["质数", "合数", "素数", "分解质因数"],

    # ===== 1.3 数的运算 =====
    "加法与乘法运算定律（交换律、结合律）": ["交换律", "结合律", "运算定律"],
    "乘法分配律": ["分配律", "乘法分配"],
    "简便计算综合": ["简便计算", "简算", "怎样简便", "能简算"],

    # ===== 1.4 式与方程 =====
    "用字母表示数": ["用字母表示", "字母表示数"],
    "含字母的式子比较大小": ["比较大小", "含字母"],
    "解方程": ["解方程", "方程", "解下列方程", "求x"],

    # ===== 1.5 比和比例 =====
    "百分数的概念": ["百分数", "百分之", "百分比", "发芽率", "出勤率", "合格率", "成活率"],
    "浓度问题": ["浓度", "盐水", "糖水", "溶液", "含盐", "含糖"],
    "比和比例的基本概念": ["比值", "化简比", "比例", "最简比", "内项", "外项"],
    "正比例和反比例": ["正比例", "反比例", "成正比", "成反比"],

    # ===== 1.6 计量单位 =====
    "长度、面积、体积单位换算": ["平方米", "平方厘米", "立方", "公顷", "换算", "单位换算"],
    "时间单位换算": ["分钟", "小时", "秒", "时间换算"],

    # ===== 2.1 平面图形 =====
    "长方形与正方形": ["长方形", "正方形", "长和宽", "边长"],
    "平行四边形与梯形": ["平行四边形", "梯形", "上底", "下底"],
    "三角形": ["三角形", "三角", "等腰", "等边", "锐角三角", "直角三角"],
    "圆": ["圆的面积", "圆的周长", "半径", "直径", "圆周率"],

    # ===== 2.2 立体图形 =====
    "长方体与正方体": ["长方体", "正方体", "棱长", "表面积和体积"],
    "圆柱": ["圆柱", "圆柱体", "侧面积", "圆柱的体积"],
    "圆锥": ["圆锥", "圆锥体", "圆锥的体积"],

    # ===== 2.3 图形运动与位置 =====
    "对称、平移与旋转": ["对称", "平移", "旋转", "轴对称"],
    "比例尺": ["比例尺", "地图", "实际距离", "图上距离"],
    "位置与方向": ["方向", "东偏北", "西偏南", "位置"],

    # ===== 2.4 面积与体积 =====
    "组合图形面积": ["组合图形", "阴影部分", "阴影面积", "不规则图形"],
    "表面积应用": ["表面积", "需要多少材料", "粉刷", "包装"],
    "体积与容积应用": ["体积", "容积", "升", "毫升", "装水", "容器"],

    # ===== 3.1 统计图 =====
    "条形统计图": ["条形统计图", "条形图"],
    "折线统计图": ["折线统计图", "折线图"],
    "扇形统计图": ["扇形统计图", "扇形图", "各部分占"],
    "平均数": ["平均数", "平均成绩", "平均分"],

    # ===== 3.2 可能性 =====
    "事件的确定性与不确定性": ["一定", "不可能", "可能", "确定事件"],
    "可能性的大小": ["可能性", "概率", "摸球", "摸到"],

    # ===== 4.1 分数乘除法应用 =====
    "分数乘除法基础题": ["分数乘", "分数除", "的几分之几是"],
    "复杂的求单位1问题": ["单位1", "单位\"1\"", "谁是单位"],

    # ===== 4.2 百分数问题 =====
    "增长与减少百分之几": ["增加百分之", "减少百分之", "多百分之", "少百分之", "增长", "降价", "涨价"],
    "折扣问题": ["打折", "折扣", "几折", "八折", "七折", "五折", "打七五折"],
    "利率税收问题": ["利率", "利息", "本金", "税", "纳税", "增值税", "年利率"],

    # ===== 4.3 工程问题 =====
    "工程问题基础题型": ["工程", "单独做", "合作", "合做", "效率", "单独完成"],
    "工程问题复杂题型": ["先做.*再做", "合做.*天.*单独", "甲做.*乙做.*剩下"],

    # ===== 4.4 分段计费 =====
    "水费的分段计费问题": ["阶梯", "分段计费", "水费", "电费", "超出部分"],

    # ===== 4.5 行程问题 =====
    "行程问题基础（路程=速度×时间）": ["速度", "路程", "时间", "行驶", "行了", "千米/时", "每小时行"],
    "相遇问题": ["相遇", "相向而行", "相向", "同时出发.*相遇"],
    "按比例解决行程问题": ["快车.*慢车", "速度比", "路程比"],

    # ===== 4.6 比与比例应用 =====
    "利用比例尺解决问题": ["比例尺", "图上.*实际", "实际.*图上"],
    "用比例解决问题": ["用比例解", "列比例", "比例方法"],
    "按比分配问题": ["按比分配", "比分配", "按.*比.*分", "长宽高的比"],

    # ===== 4.7 图象问题 =====
    "读图与图象分析": ["读图", "图象", "图表分析", "看图回答"],
}

# ---------------------------------------------------------------------------
# 难度评估关键词
# ---------------------------------------------------------------------------
DIFFICULTY_HARD_KEYWORDS = [
    "至少", "最多", "最少", "几种方案", "规律", "推理", "证明",
    "综合应用", "拓展", "思维", "挑战", "营养餐", "蛋白质含量",
    "分段计费", "阶梯计费", "牙膏.*出口", "熔成", "旋转一周",
]
DIFFICULTY_MEDIUM_KEYWORDS = [
    "如果", "已知.*求", "列式", "计算", "解方程", "应用题",
    "几分之几", "百分之几", "多少元", "多少千米", "多少天",
    "节约", "增加", "减少", "利润", "打折",
]


# ---------------------------------------------------------------------------
# 题目解析器
# ---------------------------------------------------------------------------
def parse_questions_from_text(text: str, source: str = "") -> list[dict]:
    """
    从 OCR 或 PDF 提取的文本中解析出独立的数学题目。
    
    Returns:
        list of dict, 每个含 question, source, raw_text 等
    """
    questions = []
    
    # 按"=== 第N页 ===" 分页
    pages = re.split(r"===\s*第\d+页\s*===", text)
    
    for page_text in pages:
        page_text = page_text.strip()
        if not page_text:
            continue
        
        # 尝试按题号分割: 1. 2. 3. 或 1、 2、 3、 或 (1) (2)
        # 匹配主题号 (非子题号)
        segments = re.split(
            r"(?:^|\n)\s*(\d{1,2})\s*[.、．.]\s*",
            page_text
        )
        
        # segments 格式: [前文, 题号1, 题文1, 题号2, 题文2, ...]
        if len(segments) >= 3:
            for i in range(1, len(segments) - 1, 2):
                q_num = segments[i]
                q_text = segments[i + 1].strip()
                if len(q_text) < 8:  # 太短的不是完整题目
                    continue
                    
                # 清理文本
                q_text = _clean_question_text(q_text)
                if not q_text:
                    continue
                
                questions.append({
                    "raw_num": q_num,
                    "question": q_text,
                    "source": source,
                })
        else:
            # 整段可能就是一道大题（如解答题），也解析
            cleaned = _clean_question_text(page_text)
            if cleaned and len(cleaned) > 15:
                questions.append({
                    "raw_num": "0",
                    "question": cleaned,
                    "source": source,
                })
    
    return questions


def _clean_question_text(text: str) -> str:
    """清理 OCR 识别的噪声"""
    # 去除页码
    text = re.sub(r"\b\d{2,3}\s*$", "", text, flags=re.MULTILINE)
    # 合并被换行拆分的句子
    text = re.sub(r"(?<=[^\n。？！\)）])\n(?=[^\n\d（(=])", "", text)
    # 去除多余空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去除首尾空白
    text = text.strip()
    return text


# ---------------------------------------------------------------------------
# 知识点分类器
# ---------------------------------------------------------------------------
def classify_topic(question_text: str) -> tuple[str, float]:
    """
    根据关键词将题目分类到最匹配的知识点。
    
    Returns:
        (topic_name, confidence_score)
        confidence_score 0~1, 越高越确定
    """
    scores: dict[str, int] = {}
    text_lower = question_text.lower()
    
    for topic, keywords in TOPIC_KEYWORDS.items():
        score = 0
        for kw in keywords:
            if re.search(kw, text_lower):
                score += 1
        if score > 0:
            scores[topic] = score
    
    if not scores:
        return ("", 0.0)
    
    # 取得分最高的
    best_topic = max(scores, key=scores.get)
    best_score = scores[best_topic]
    max_possible = len(TOPIC_KEYWORDS.get(best_topic, []))
    confidence = min(best_score / max(max_possible, 1), 1.0)
    
    return (best_topic, confidence)


# ---------------------------------------------------------------------------
# 难度评估器
# ---------------------------------------------------------------------------
def assess_difficulty(question_text: str) -> int:
    """
    基于关键词和题目特征评估难度 (1=基础, 2=中等, 3=困难)
    
    评估标准:
    - 文本长度：越长通常越复杂
    - 子问题数量：含 (1)(2)(3) 等子问题
    - 关键词匹配
    - 数学运算复杂度
    """
    score = 0
    
    # 1. 文本长度
    length = len(question_text)
    if length > 200:
        score += 2
    elif length > 100:
        score += 1
    
    # 2. 子问题数量
    sub_questions = len(re.findall(r"[（(]\s*\d+\s*[)）]", question_text))
    if sub_questions >= 3:
        score += 2
    elif sub_questions >= 2:
        score += 1
    
    # 3. 困难关键词
    for kw in DIFFICULTY_HARD_KEYWORDS:
        if re.search(kw, question_text):
            score += 2
            break
    
    # 4. 中等关键词
    for kw in DIFFICULTY_MEDIUM_KEYWORDS:
        if re.search(kw, question_text):
            score += 1
            break
    
    # 5. 包含多种运算
    ops = len(set(re.findall(r"[+\-×÷*/]", question_text)))
    if ops >= 3:
        score += 1
    
    # 映射到 1-3
    if score >= 5:
        return 3
    elif score >= 2:
        return 2
    else:
        return 1


# ---------------------------------------------------------------------------
# 题型推断
# ---------------------------------------------------------------------------
def infer_question_type(question_text: str) -> str:
    """推断题目类型: 选择/填空/判断/计算/解答"""
    text = question_text[:100]  # 只看开头部分
    
    if re.search(r"[ABCD][.、．]|[（(]\s*[)）]\s*\n\s*A", question_text):
        return "选择"
    if "判断" in text or "对还是错" in text or "是否正确" in text:
        return "判断"
    if re.search(r"[（(]\s*[)）]", text) and "=" not in text[:30]:
        return "填空"
    if "解方程" in text or "计算" in text or "简算" in text or "直接写出得数" in text:
        return "计算"
    return "解答"


# ---------------------------------------------------------------------------
# 生成唯一 ID
# ---------------------------------------------------------------------------
def generate_id(question_text: str, source: str = "") -> str:
    """基于内容哈希生成唯一 ID，格式: IMP-XXXX"""
    content = f"{source}:{question_text}"
    h = hashlib.md5(content.encode("utf-8")).hexdigest()[:6].upper()
    return f"IMP-{h}"


# ---------------------------------------------------------------------------
# 主导入流程
# ---------------------------------------------------------------------------
def import_from_text_file(filepath: str, auto_save: bool = True) -> list[dict]:
    """
    从文本文件自动导入题目到题库。
    
    流程: 读取文本 → 解析题目 → 分类知识点 → 评估难度 → 保存到 JSON
    
    Args:
        filepath: 文本文件路径 (OCR 输出或 PDF 提取文本)
        auto_save: 是否自动保存到 imported_exercises.json
    
    Returns:
        成功导入的题目列表
    """
    source = os.path.basename(filepath)
    
    with open(filepath, "r", encoding="utf-8") as f:
        text = f.read()
    
    print(f"\n[IMPORT] {source}")
    print(f"  Text length: {len(text)} chars")
    
    # 1. 解析题目
    raw_questions = parse_questions_from_text(text, source)
    print(f"  Parsed: {len(raw_questions)} raw questions")
    
    # 2. 分类、评估、结构化
    imported = []
    classified_count = 0
    unclassified = []
    
    for rq in raw_questions:
        q_text = rq["question"]
        
        # 分类
        topic, confidence = classify_topic(q_text)
        if not topic or confidence < 0.1:
            unclassified.append(q_text[:60])
            continue
        
        classified_count += 1
        
        # 评估难度
        difficulty = assess_difficulty(q_text)
        
        # 推断题型
        q_type = infer_question_type(q_text)
        
        # 生成 ID
        q_id = generate_id(q_text, source)
        
        imported.append({
            "id": q_id,
            "type": q_type,
            "difficulty": difficulty,
            "question": q_text,
            "answer": "(待补充)",
            "explanation": "(待补充)",
            "source": source,
            "topic": topic,
            "confidence": round(confidence, 2),
        })
    
    print(f"  Classified: {classified_count} questions")
    if unclassified:
        print(f"  Unclassified: {len(unclassified)} (skipped)")
        for u in unclassified[:3]:
            print(f"    - {u}...")
    
    # 3. 保存
    if auto_save and imported:
        _save_imported(imported)
    
    return imported


def import_from_pdf(pdf_path: str, use_ocr: bool = False) -> list[dict]:
    """
    从 PDF 文件导入题目。
    
    Args:
        pdf_path: PDF 文件路径
        use_ocr: 是否使用 OCR (扫描件设为 True)
    """
    import fitz
    
    pdf_name = os.path.basename(pdf_path)
    txt_name = os.path.splitext(pdf_name)[0]
    
    if use_ocr:
        # 先 OCR 提取
        from ocr_extract import extract_pdf_with_ocr
        txt_path = os.path.join(DATA_DIR, f"{txt_name}_ocr.txt")
        if not os.path.exists(txt_path):
            print(f"[OCR] Extracting {pdf_name}...")
            extract_pdf_with_ocr(pdf_path, txt_path)
    else:
        # 直接文本提取
        txt_path = os.path.join(DATA_DIR, f"{txt_name}.txt")
        if not os.path.exists(txt_path):
            doc = fitz.open(pdf_path)
            pages_text = []
            for i, page in enumerate(doc):
                pages_text.append(f"=== 第{i+1}页 ===\n{page.get_text()}")
            doc.close()
            os.makedirs(DATA_DIR, exist_ok=True)
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("\n\n".join(pages_text))
    
    return import_from_text_file(txt_path)


def import_from_docx(docx_path: str) -> list[dict]:
    """
    从 Word (.docx) 文件导入题目。

    提取所有段落和表格中的文本，保存为 txt 后走标准导入管线。

    Args:
        docx_path: Word 文档路径
    """
    from docx import Document

    docx_name = os.path.basename(docx_path)
    txt_name = os.path.splitext(docx_name)[0]
    txt_path = os.path.join(DATA_DIR, f"{txt_name}.txt")

    print(f"[DOCX] Extracting text from {docx_name}...")
    doc = Document(docx_path)

    parts = []

    # 提取所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取表格中的文本（试卷常把题目放在表格里）
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append("  ".join(row_texts))

    if not parts:
        print(f"  [WARN] No text found in {docx_name}")
        return []

    full_text = "\n".join(parts)
    print(f"  Extracted {len(parts)} paragraphs, {len(full_text)} chars")

    # 保存为 txt
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(full_text)
    print(f"  Saved to {txt_path}")

    # 走标准导入管线
    return import_from_text_file(txt_path)


def _save_imported(new_exercises: list[dict]):
    """保存导入结果到 JSON, 与已有数据合并去重"""
    os.makedirs(DATA_DIR, exist_ok=True)
    
    # 加载已有数据
    existing = {}
    if os.path.exists(IMPORT_FILE):
        try:
            with open(IMPORT_FILE, "r", encoding="utf-8") as f:
                existing = json.load(f)
        except (json.JSONDecodeError, OSError):
            existing = {}
    
    existing_ids = set()
    for exs in existing.values():
        for ex in exs:
            existing_ids.add(ex.get("id"))
    
    # 合并新题目
    added_count = 0
    for ex in new_exercises:
        topic = ex.pop("topic")  # topic 是分类用的，不存在题目本身
        conf = ex.pop("confidence", 0)  # confidence 也不存题目本身
        
        if ex["id"] in existing_ids:
            continue
        
        if topic not in existing:
            existing[topic] = []
        existing[topic].append(ex)
        existing_ids.add(ex["id"])
        added_count += 1
    
    # 写入
    with open(IMPORT_FILE, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)
    
    total = sum(len(v) for v in existing.values())
    print(f"  Saved: +{added_count} new (total {total} in imported_exercises.json)")


# ---------------------------------------------------------------------------
# 批量扫描 data/ 目录
# ---------------------------------------------------------------------------
def scan_and_import_all():
    """扫描 data/ 目录下所有文本文件，自动导入"""
    if not os.path.isdir(DATA_DIR):
        print("[WARN] data/ directory not found")
        return
    
    text_files = [
        f for f in os.listdir(DATA_DIR)
        if f.endswith((".txt",)) and not f.startswith(".")
    ]
    
    if not text_files:
        print("[INFO] No text files found in data/")
        return
    
    print(f"[SCAN] Found {len(text_files)} text files in data/")
    all_imported = []
    
    for tf in sorted(text_files):
        filepath = os.path.join(DATA_DIR, tf)
        imported = import_from_text_file(filepath, auto_save=True)
        all_imported.extend(imported)
    
    print(f"\n{'='*50}")
    print(f"[DONE] Total imported: {len(all_imported)} questions")
    
    # 统计分布
    from collections import Counter
    topic_dist = Counter(ex.get("topic", "?") for ex in all_imported)
    diff_dist = Counter(ex.get("difficulty", 0) for ex in all_imported)
    type_dist = Counter(ex.get("type", "?") for ex in all_imported)
    
    print(f"\n  By difficulty: {dict(sorted(diff_dist.items()))}")
    print(f"  By type: {dict(type_dist)}")
    print(f"  Top topics:")
    # topic was popped in _save_imported, so recalculate from file
    if os.path.exists(IMPORT_FILE):
        with open(IMPORT_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        for topic, exs in sorted(data.items(), key=lambda x: -len(x[1]))[:10]:
            print(f"    {topic}: {len(exs)} questions")


# ---------------------------------------------------------------------------
# 命令行入口
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    args = sys.argv[1:]
    
    if not args:
        # 默认: 扫描 data/ 目录
        scan_and_import_all()
    elif args[0] == "--from-pdf":
        # 从 PDF 导入
        if len(args) < 2:
            print("Usage: python import_questions.py --from-pdf <pdf_path> [--ocr]")
            sys.exit(1)
        pdf_path = args[1]
        if not os.path.isabs(pdf_path):
            pdf_path = os.path.join(BASE_DIR, pdf_path)
        use_ocr = "--ocr" in args
        import_from_pdf(pdf_path, use_ocr=use_ocr)
    elif args[0] == "--from-docx":
        # 从 Word 文档导入
        if len(args) < 2:
            print("Usage: python import_questions.py --from-docx <docx_path>")
            sys.exit(1)
        docx_path = args[1]
        if not os.path.isabs(docx_path):
            docx_path = os.path.join(BASE_DIR, docx_path)
        import_from_docx(docx_path)
    else:
        # 导入指定文本文件，也支持直接传 .docx 文件
        for filepath in args:
            if not os.path.isabs(filepath):
                filepath = os.path.join(BASE_DIR, filepath)
            if filepath.lower().endswith(".docx"):
                import_from_docx(filepath)
            else:
                import_from_text_file(filepath)
